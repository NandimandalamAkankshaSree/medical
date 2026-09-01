import os
import re
import json
import base64
import difflib
import urllib.request
import urllib.parse
import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
import pypdf
from app.core.config import settings

try:
    import docx
except ImportError:
    docx = None


class DocumentParser:
    """
    Extracts text, metadata, laboratory tables, and clinical sections
    from medical documents (PDFs, DOCX files, Images, Scanned reports) while preserving
    page numbers, sections, and exact coordinates.
    """

    @staticmethod
    def clean_text(s: str) -> str:
        if not s:
            return ""
        return (
            s.replace("\ufffd", "-")
            .replace("\u2013", "-")
            .replace("\u2014", "-")
            .replace("–", "-")
            .replace("—", "-")
            .strip()
        )

    @staticmethod
    def extract_text_with_pages(file_path: str) -> List[Dict[str, Any]]:
        """
        Extracts text page by page (or section by section) with page numbers.
        Supports PDF, DOCX, and Image formats.
        """
        pages = []
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            try:
                reader = pypdf.PdfReader(file_path)
                for idx, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    pages.append({
                        "page_number": idx + 1,
                        "text": DocumentParser.clean_text(text),
                        "char_count": len(text)
                    })
            except Exception as e:
                pages.append({
                    "page_number": 1,
                    "text": f"Error reading PDF: {str(e)}",
                    "char_count": 0
                })

        elif ext in [".docx", ".doc"]:
            try:
                if docx is not None:
                    doc = docx.Document(file_path)
                    lines = []
                    for p in doc.paragraphs:
                        if p.text.strip():
                            lines.append(DocumentParser.clean_text(p.text))
                    
                    # Also append table content as structured text
                    for t in doc.tables:
                        for row in t.rows:
                            cells = [DocumentParser.clean_text(c.text) for c in row.cells]
                            if any(cells):
                                lines.append(" | ".join(cells))

                    text = "\n".join(lines)
                    pages.append({
                        "page_number": 1,
                        "text": text,
                        "char_count": len(text)
                    })
                else:
                    # Fallback zip reading if docx package unavailable
                    import zipfile
                    import xml.etree.ElementTree as ET
                    with zipfile.ZipFile(file_path) as z:
                        xml_content = z.read("word/document.xml")
                        tree = ET.fromstring(xml_content)
                        text = " ".join([node.text for node in tree.iter() if node.text])
                        pages.append({
                            "page_number": 1,
                            "text": DocumentParser.clean_text(text),
                            "char_count": len(text)
                        })
            except Exception as e:
                pages.append({
                    "page_number": 1,
                    "text": f"Error reading DOCX: {str(e)}",
                    "char_count": 0
                })

        elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
            # 1. Try Gemini Vision multimodal OCR if key is configured
            gemini_text = DocumentParser.extract_image_with_gemini(file_path)
            if gemini_text:
                pages.append({
                    "page_number": 1,
                    "text": DocumentParser.clean_text(gemini_text),
                    "char_count": len(gemini_text)
                })
            else:
                # 2. Try pytesseract OCR if available
                try:
                    import pytesseract
                    from PIL import Image
                    img = Image.open(file_path)
                    text = pytesseract.image_to_string(img)
                    pages.append({
                        "page_number": 1,
                        "text": DocumentParser.clean_text(text),
                        "char_count": len(text)
                    })
                except Exception:
                    pages.append({
                        "page_number": 1,
                        "text": "",
                        "char_count": 0
                    })
        else:
            # If it's a plain text or unknown file
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                    pages.append({
                        "page_number": 1,
                        "text": DocumentParser.clean_text(text),
                        "char_count": len(text)
                    })
            except Exception:
                pages.append({
                    "page_number": 1,
                    "text": "",
                    "char_count": 0
                })

        return pages

    @staticmethod
    def extract_image_with_gemini(file_path: str) -> Optional[str]:
        """
        Uses Gemini Vision API via HTTPS to extract text from medical images or detect non-medical images.
        """
        api_key = settings.GEMINI_API_KEY.strip()
        if not api_key:
            return None
        
        try:
            with open(file_path, "rb") as f:
                img_bytes = f.read()
            b64_img = base64.b64encode(img_bytes).decode("utf-8")
            
            ext = os.path.splitext(file_path)[1].lower().replace(".", "")
            mime_type = f"image/{ext}" if ext in ["png", "jpeg", "jpg", "webp"] else "image/jpeg"
            if ext == "jpg":
                mime_type = "image/jpeg"

            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
            payload = {
                "contents": [
                    {
                        "parts": [
                            {
                                "inline_data": {
                                    "mime_type": mime_type,
                                    "data": b64_img
                                }
                            },
                            {
                                "text": "Extract all clinical text, patient names, dates, test tables, and laboratory parameters from this medical image. If this image is NOT a medical document, medical report, prescription, or lab result, reply only with: [NOT_A_MEDICAL_DOCUMENT]"
                            }
                        ]
                    }
                ],
                "generationConfig": {"temperature": 0.0, "maxOutputTokens": 2048}
            }
            req = urllib.request.Request(
                url,
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type": "application/json"},
                method="POST"
            )
            with urllib.request.urlopen(req, timeout=10) as res:
                if res.status == 200:
                    resp = json.loads(res.read().decode("utf-8"))
                    candidates = resp.get("candidates", [])
                    if candidates:
                        text = candidates[0].get("content", {}).get("parts", [{}])[0].get("text", "")
                        if "[NOT_A_MEDICAL_DOCUMENT]" in text:
                            return ""
                        return text
        except Exception as e:
            print(f"Gemini Vision OCR notice: {e}")
        return None

    @staticmethod
    def validate_medical_document(parsed: Dict[str, Any], full_text: str) -> Tuple[bool, str]:
        """
        Validates whether the uploaded file or image is a genuine medical report or healthcare document.
        Returns (is_valid: bool, error_message: str).
        """
        # 1. If lab parameters were extracted
        if parsed.get("lab_parameters") and len(parsed["lab_parameters"]) > 0:
            return True, "Valid Medical Laboratory Report"

        text_lower = (full_text or "").lower().strip()
        
        # If text is virtually empty (e.g. non-text photo or unreadable image)
        if len(text_lower) < 20:
            return False, "The uploaded file does not contain readable clinical text or medical laboratory parameters. Please upload a clear medical report or lab result."

        # 2. Check for medical keywords and clinical markers
        medical_score = 0
        strong_medical_indicators = [
            "patient", "doctor", "dr.", "hospital", "clinic", "laboratory", "diagnostic", "diagnostics",
            "reference range", "normal range", "specimen", "investigation", "findings", "impression",
            "prescription", "dosage", "tablet", "capsule", "rx", "serum", "blood", "urine",
            "creatinine", "hemoglobin", "glucose", "cholesterol", "potassium", "bilirubin", "platelet",
            "wbc", "rbc", "egfr", "bun", "urea", "uric acid", "tsh", "alt", "ast", "sgpt", "sgot",
            "ultrasound", "x-ray", "ct scan", "mri", "biopsy", "pathology", "radiology", "cardiology",
            "nephrology", "endocrinology", "ecg", "ekg", "symptoms", "diagnosis", "clinical",
            "mg/dl", "meq/l", "mmol/l", "g/dl", "bpm", "mmhg"
        ]

        for word in strong_medical_indicators:
            if word in text_lower:
                medical_score += 1

        # A valid medical document must have at least 2 distinct clinical indicators
        if medical_score >= 2:
            return True, "Valid Medical Document"

        return False, "The uploaded file does not contain recognizable medical or healthcare information. Please upload a correct medical report (e.g., Blood Test, Diagnostic Lab Report, Doctor Prescription, or Scan)."

    @staticmethod
    def normalize_name_tokens(name: str) -> List[str]:
        if not name:
            return []
        # Remove common titles and prefixes
        cleaned = re.sub(
            r"\b(mr|mrs|ms|miss|dr|doctor|prof|shri|smt|kumari|kumar|master|baby\s+of|pt|patient|patient\s+name|self|user)\b\.?",
            "",
            name,
            flags=re.IGNORECASE
        )
        tokens = [t.lower() for t in re.findall(r"[A-Za-z]+", cleaned)]
        return [t for t in tokens if len(t) >= 2]

    @staticmethod
    def is_patient_name_match(report_patient_name: str, logged_in_patient_name: str) -> Tuple[bool, str]:
        """
        Validates if the patient name extracted from the uploaded report matches the logged in patient profile.
        Returns (is_match: bool, reason: str).
        """
        if not report_patient_name:
            return True, "No explicit patient name found in report."

        rep_clean = report_patient_name.strip().lower()
        generic_names = {"patient", "unknown", "self", "user", "client", "n/a", "na", "sample", "anonymous", "null", "none"}
        if rep_clean in generic_names or not rep_clean:
            return True, "Generic patient name permitted."

        if not logged_in_patient_name or logged_in_patient_name.strip().lower() in generic_names:
            return True, "Default account profile permitted."

        rep_tokens = DocumentParser.normalize_name_tokens(report_patient_name)
        login_tokens = DocumentParser.normalize_name_tokens(logged_in_patient_name)

        if not rep_tokens or not login_tokens:
            return True, "Insufficient name tokens to determine mismatch."

        # Check for significant token overlap (e.g. 'akanksha', 'sree', 'nandimandalam')
        overlap = set(rep_tokens).intersection(set(login_tokens))
        if overlap:
            return True, f"Matching name token(s): {', '.join(overlap)}"

        # Check fuzzy sequence similarity
        rep_str = " ".join(rep_tokens)
        login_str = " ".join(login_tokens)
        ratio = difflib.SequenceMatcher(None, rep_str, login_str).ratio()
        if ratio >= 0.50:
            return True, f"Fuzzy name match ({ratio:.2f})"

        # Definite mismatch with another person
        return False, f"Document belongs to '{report_patient_name.strip()}', which does not match your logged-in account name '{logged_in_patient_name.strip()}'."

    @staticmethod
    def parse_medical_report(file_path: str) -> Dict[str, Any]:
        """
        Parses a medical report (PDF, DOCX, Image) into structured parameters, patient info,
        hospital details, doctor details, and laboratory findings.
        """
        pages = DocumentParser.extract_text_with_pages(file_path)
        full_text = "\n".join([p["text"] for p in pages])
        ext = os.path.splitext(file_path)[1].lower()

        # Extract Document Title / Type
        doc_type = "Medical Report"
        first_lines = [line.strip() for line in full_text.splitlines() if line.strip()]
        if first_lines:
            first_line = first_lines[0].lower()
            if "blood report" in first_line or "complete blood count" in full_text.lower() or "cbc" in full_text.lower():
                doc_type = "Blood Report"
            elif "diabetes report" in first_line or "hba1c" in full_text.lower() or "glucose" in full_text.lower():
                doc_type = "Diabetes Report"
            elif "lipid profile" in first_line or "cholesterol" in full_text.lower():
                doc_type = "Lipid Profile"
            elif "thyroid report" in first_line or "tsh" in full_text.lower():
                doc_type = "Thyroid Report"
            elif "kidney" in full_text.lower() or "renal" in full_text.lower() or "creatinine" in full_text.lower() or "nephrology" in full_text.lower():
                doc_type = "Kidney Report"
            elif "prescription" in first_line or "rx" in first_line:
                doc_type = "Prescription"
            else:
                doc_type = first_lines[0][:50]

        # Extract Patient Details
        patient_id_match = re.search(r"Patient ID:\s*([\w\-]+)", full_text, re.IGNORECASE)
        name_match = re.search(
            r"(?:Patient\s+Name|Name\s+of\s+Patient|Patient|Pt\.?\s*Name|Client\s+Name|Name):\s*([A-Za-z\.\s]+?)(?=(?:\s+Age|\s+Sex|\s+Gender|\s+ID|\s+Ref|\s+Dr\.|\s+Date|\s+Hospital|\n|\r|,|$))",
            full_text,
            re.IGNORECASE
        )
        age_match = re.search(r"Age(?:/Sex)?:\s*(\d+)", full_text, re.IGNORECASE)
        sex_match = re.search(r"(?:Sex|Gender):\s*([A-Za-z]+)", full_text, re.IGNORECASE)
        if not sex_match:
            sex_combo_match = re.search(r"Age/Sex:\s*\d+\s*(?:years)?\s*/\s*([A-Za-z]+)", full_text, re.IGNORECASE)
            if sex_combo_match:
                sex_match = sex_combo_match

        hospital_match = re.search(r"Hospital:\s*([^\n\r]+)", full_text, re.IGNORECASE)
        doctor_match = re.search(r"(?:Doctor|Physician|Dr\.):\s*([^\n\r]+)", full_text, re.IGNORECASE)
        
        # Date match
        report_date = DocumentParser.extract_date(full_text)

        findings_match = re.search(r"(?:Clinical Impression|Findings)\s*([\s\S]*?)(?=(Doctor Observations|Recommendations|Important Note|This document|$))", full_text, re.IGNORECASE)
        doctor_obs_match = re.search(r"Doctor Observations\s*([\s\S]*?)(?=(Recommendations|Important Note|This document|$))", full_text, re.IGNORECASE)
        recommendations_match = re.search(r"Recommendations\s*([\s\S]*?)(?=(Important Note|This document|$))", full_text, re.IGNORECASE)

        patient_id = patient_id_match.group(1).strip() if patient_id_match else None
        
        # Clean patient name
        if name_match:
            raw_name = name_match.group(1).strip()
            cleaned_name = re.sub(r"^(?:mr|mrs|ms|miss|dr|doctor|master|shri|smt|kumari|kumar|baby\s+of)\.?\s+", "", raw_name, flags=re.IGNORECASE).strip()
            name = cleaned_name if cleaned_name else raw_name
        else:
            name = "Patient"
        age = int(age_match.group(1)) if age_match else None
        sex = sex_match.group(1).strip() if sex_match else None
        hospital = hospital_match.group(1).strip() if hospital_match else "CityCare Multispeciality Hospital"
        doctor = doctor_match.group(1).strip() if doctor_match else "Attending Physician"

        findings = findings_match.group(1).strip() if findings_match else None
        doctor_obs = doctor_obs_match.group(1).strip() if doctor_obs_match else None
        recommendations = recommendations_match.group(1).strip() if recommendations_match else None

        # Extract Lab Parameter Values
        lab_parameters = []
        if ext in [".docx", ".doc"] and docx is not None:
            lab_parameters = DocumentParser.extract_docx_tables(file_path, doc_type)

        if not lab_parameters:
            lab_parameters = DocumentParser.extract_lab_parameters(full_text, doc_type)

        # Generate Patient Friendly Summaries
        quick_summary, detailed_summary = DocumentParser.generate_patient_summaries(
            doc_type=doc_type,
            lab_parameters=lab_parameters,
            findings=findings,
            doctor_obs=doctor_obs
        )

        return {
            "document_type": doc_type,
            "patient_id": patient_id,
            "patient_name": name,
            "age": age,
            "sex": sex,
            "hospital_name": hospital,
            "doctor_name": doctor,
            "report_date": report_date,
            "findings": findings,
            "doctor_observations": doctor_obs,
            "recommendations": recommendations,
            "lab_parameters": lab_parameters,
            "quick_summary": quick_summary,
            "detailed_summary": detailed_summary,
            "page_count": len(pages),
            "pages": pages,
            "full_text": full_text
        }

    @staticmethod
    def extract_date(text: str) -> str:
        """
        Parses report dates into standard YYYY-MM-DD format.
        """
        # Formats like "Date of Report: 24 August 2026", "Report Date: 2026-08-24", "Date: 24-Aug-2026"
        date_patterns = [
            r"(?:Date of Report|Report Date|Date|Previous Report):\s*(\d{1,2}\s+[A-Za-z]+\s+\d{4})",
            r"(?:Date of Report|Report Date|Date|Previous Report):\s*([A-Za-z]+\s+\d{1,2},?\s+\d{4})",
            r"(?:Date of Report|Report Date|Date|Previous Report):\s*(\d{4}-\d{1,2}-\d{1,2})",
            r"(?:Date of Report|Report Date|Date|Previous Report):\s*(\d{1,2}/\d{1,2}/\d{4})",
            r"(\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})"
        ]

        for pat in date_patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                d_str = m.group(1).strip()
                try:
                    # Attempt standard parsing
                    for fmt in ["%d %B %Y", "%B %d %Y", "%B %d, %Y", "%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%b-%Y"]:
                        try:
                            dt = datetime.datetime.strptime(d_str.replace(",", ""), fmt)
                            return dt.strftime("%Y-%m-%d")
                        except ValueError:
                            pass
                except Exception:
                    pass
                return d_str

        return datetime.date.today().isoformat()

    @staticmethod
    def extract_docx_tables(file_path: str, doc_type: str) -> List[Dict[str, Any]]:
        """
        Extracts structured lab parameters directly from DOCX tables.
        """
        lab_results = []
        if docx is None:
            return lab_results

        try:
            doc = docx.Document(file_path)
            for t in doc.tables:
                if len(t.rows) < 2:
                    continue

                header_cells = [DocumentParser.clean_text(c.text).lower() for c in t.rows[0].cells]
                
                # Check if this table has test/result columns
                test_idx = 0
                res_idx = 1
                ref_idx = 2 if len(header_cells) > 2 else -1

                for idx, h in enumerate(header_cells):
                    if "test" in h or "parameter" in h or "investigation" in h:
                        test_idx = idx
                    elif "result" in h or "value" in h:
                        res_idx = idx
                    elif "reference" in h or "range" in h or "interval" in h:
                        ref_idx = idx

                for row in t.rows[1:]:
                    cells = [DocumentParser.clean_text(c.text) for c in row.cells]
                    if len(cells) <= max(test_idx, res_idx):
                        continue

                    param_name = cells[test_idx].strip()
                    res_raw = cells[res_idx].strip()
                    ref_raw = cells[ref_idx].strip() if ref_idx != -1 and len(cells) > ref_idx else ""

                    if not param_name or any(h in param_name.lower() for h in ["test", "parameter", "investigation"]):
                        continue

                    num_val, res_clean, unit = DocumentParser.parse_val_unit(res_raw)
                    min_ref, max_ref = DocumentParser.parse_reference_range(ref_raw, num_val)
                    status, interpretation = DocumentParser.evaluate_status(param_name, num_val, res_clean, min_ref, max_ref, ref_raw)
                    category = DocumentParser.determine_category(param_name, doc_type)

                    lab_results.append({
                        "parameter_name": param_name,
                        "result_value": res_clean,
                        "numeric_value": num_val,
                        "unit": unit,
                        "reference_range": ref_raw or "Standard",
                        "min_ref": min_ref,
                        "max_ref": max_ref,
                        "status": status,
                        "interpretation": interpretation,
                        "category": category,
                        "page_number": 1,
                        "section_name": "Laboratory Investigations"
                    })
        except Exception as e:
            print(f"Error extracting docx tables: {e}")

        return lab_results

    @staticmethod
    def extract_lab_parameters(text: str, doc_type: str) -> List[Dict[str, Any]]:
        """
        Parses structured laboratory results with numeric values, units, and reference ranges.
        Handles multiline blocks, table pipe lines, and tabular whitespace lines.
        """
        lab_results = []
        seen_params = set()

        # Strategy 1: Multiline test blocks:
        # Parameter Name \n Result: X \n Unit: Y \n Reference Range: Z
        test_blocks = re.findall(
            r"([A-Za-z0-9\s\-\(\)\/\%]+)\n\s*Result:\s*([0-9\.]+|Negative|Positive|[A-Za-z\+\d]+)\s*(?:Unit:\s*([^\n\r]+))?\n\s*Reference Range:\s*([^\n\r]+)",
            text
        )
        for param_name, res_val, unit, ref_range in test_blocks:
            p_clean = DocumentParser.clean_text(param_name)
            if not p_clean or any(h in p_clean.lower() for h in ["patient details", "laboratory results", "findings", "synthetic"]):
                continue
            if p_clean in seen_params:
                continue

            num_val = None
            try:
                num_val = float(res_val.strip())
            except ValueError:
                pass

            min_ref, max_ref = DocumentParser.parse_reference_range(ref_range, num_val)
            status, interpretation = DocumentParser.evaluate_status(p_clean, num_val, res_val.strip(), min_ref, max_ref, ref_range.strip())
            category = DocumentParser.determine_category(p_clean, doc_type)

            lab_results.append({
                "parameter_name": p_clean,
                "result_value": res_val.strip(),
                "numeric_value": num_val,
                "unit": unit.strip() if unit else "",
                "reference_range": ref_range.strip() if ref_range else "Standard",
                "min_ref": min_ref,
                "max_ref": max_ref,
                "status": status,
                "interpretation": interpretation,
                "category": category,
                "page_number": 1,
                "section_name": "Laboratory Results"
            })
            seen_params.add(p_clean)

        # Strategy 2: Pipe separated or tabular text lines (e.g. from DOCX/PDF table extractions)
        for line in text.splitlines():
            line_clean = DocumentParser.clean_text(line)
            if " | " in line_clean:
                parts = [p.strip() for p in line_clean.split(" | ") if p.strip()]
                if len(parts) >= 2:
                    p_name = parts[0]
                    if any(h in p_name.lower() for h in ["test", "parameter", "investigation", "hospital", "patient"]):
                        continue
                    if p_name in seen_params:
                        continue

                    res_raw = parts[1]
                    ref_raw = parts[2] if len(parts) > 2 else ""

                    num_val, res_val, unit = DocumentParser.parse_val_unit(res_raw)
                    min_ref, max_ref = DocumentParser.parse_reference_range(ref_raw, num_val)
                    status, interpretation = DocumentParser.evaluate_status(p_name, num_val, res_val, min_ref, max_ref, ref_raw)
                    category = DocumentParser.determine_category(p_name, doc_type)

                    lab_results.append({
                        "parameter_name": p_name,
                        "result_value": res_val,
                        "numeric_value": num_val,
                        "unit": unit,
                        "reference_range": ref_raw or "Standard",
                        "min_ref": min_ref,
                        "max_ref": max_ref,
                        "status": status,
                        "interpretation": interpretation,
                        "category": category,
                        "page_number": 1,
                        "section_name": "Laboratory Results"
                    })
                    seen_params.add(p_name)

        return lab_results

    @staticmethod
    def parse_val_unit(res_str: str) -> Tuple[Optional[float], str, str]:
        """
        Extracts numerical value, clean string representation, and unit from result string.
        """
        res_str = DocumentParser.clean_text(res_str)
        if not res_str:
            return None, "", ""

        # Range result like "2-4 / HPF" or "4-6 / HPF"
        m_range = re.match(r"^([0-9\.]+)\s*[\-\–]\s*([0-9\.]+)\s*(.*)$", res_str)
        if m_range:
            avg_val = round((float(m_range.group(1)) + float(m_range.group(2))) / 2.0, 2)
            unit = m_range.group(3).strip()
            return avg_val, res_str, unit

        # Single number e.g. "2.1 mg/dL", "138 mmol/L", "10.8 g/dL"
        m_single = re.match(r"^([0-9\.]+)\s*(.*)$", res_str)
        if m_single:
            try:
                num = float(m_single.group(1))
                unit = m_single.group(2).strip()
                return num, res_str, unit
            except ValueError:
                pass

        # Qualitative value like "+2", "+3", "Negative", "Positive"
        return None, res_str, ""

    @staticmethod
    def parse_reference_range(ref_str: str, numeric_val: Optional[float]) -> Tuple[Optional[float], Optional[float]]:
        """
        Parses reference intervals like "13.0–17.0", "0.7–1.3 mg/dL", "<140", ">90", "3.5–5.0", etc.
        """
        if not ref_str:
            return None, None

        cleaned = DocumentParser.clean_text(ref_str)

        # Range format: "X - Y"
        range_match = re.search(r"([0-9\.]+)\s*[\-\–]\s*([0-9\.]+)", cleaned)
        if range_match:
            try:
                return float(range_match.group(1)), float(range_match.group(2))
            except ValueError:
                pass

        # Greater than format: "> X" or ">= X"
        gt_match = re.search(r">=?\s*([0-9\.]+)", cleaned)
        if gt_match:
            try:
                return float(gt_match.group(1)), 99999.0
            except ValueError:
                pass

        # Less than format: "< X" or "<= X"
        lt_match = re.search(r"<=?\s*([0-9\.]+)", cleaned)
        if lt_match:
            try:
                return 0.0, float(lt_match.group(1))
            except ValueError:
                pass

        return None, None

    @staticmethod
    def evaluate_status(
        name: str,
        num_val: Optional[float],
        raw_val: str,
        min_ref: Optional[float],
        max_ref: Optional[float],
        ref_range: str
    ) -> Tuple[str, str]:
        """
        Evaluates clinical status (LOW, NORMAL, HIGH) and provides simple patient explanation.
        """
        if raw_val.lower() in ["negative", "normal", "nil", "none"]:
            return "NORMAL", "Result is negative / normal, which is within expected healthy limits."
        if raw_val.lower() in ["positive", "+1", "+2", "+3", "+4"]:
            return "HIGH", f"Result indicates presence of marker ({raw_val}), flagged for medical review."

        if num_val is None or (min_ref is None and max_ref is None):
            return "NORMAL", "Result recorded from medical document."

        if min_ref is not None and num_val < min_ref:
            return "LOW", f"Your {name} level ({num_val}) is below the standard reference minimum ({min_ref})."
        elif max_ref is not None and num_val > max_ref:
            return "HIGH", f"Your {name} level ({num_val}) is elevated above the standard reference maximum ({max_ref})."
        else:
            return "NORMAL", f"Your {name} level ({num_val}) is within the standard reference range ({ref_range})."

    @staticmethod
    def determine_category(name: str, doc_type: str) -> str:
        name_l = name.lower()
        if any(w in name_l for w in ["creatinine", "bun", "urea", "egfr", "uric acid", "kidney", "renal"]):
            return "Renal Function"
        elif any(w in name_l for w in ["sodium", "potassium", "chloride", "bicarbonate", "electrolyte"]):
            return "Electrolytes"
        elif any(w in name_l for w in ["urine", "protein", "rbc", "wbc", "casts", "epithelial", "hpf"]):
            return "Urinalysis"
        elif any(w in name_l for w in ["hemoglobin", "wbc", "rbc", "platelet", "hematocrit", "mcv", "mch"]):
            return "Hematology"
        elif any(w in name_l for w in ["glucose", "hba1c", "insulin"]):
            return "Diabetes / Glycemic"
        elif any(w in name_l for w in ["cholesterol", "triglyceride", "ldl", "hdl", "vldl"]):
            return "Lipid Profile"
        elif any(w in name_l for w in ["tsh", "t3", "t4", "thyroid"]):
            return "Thyroid Function"
        elif any(w in name_l for w in ["ast", "alt", "bilirubin", "sgot", "sgpt", "albumin"]):
            return "Liver Function"
        return "General Diagnostics"

    @staticmethod
    def generate_patient_summaries(
        doc_type: str,
        lab_parameters: List[Dict[str, Any]],
        findings: Optional[str],
        doctor_obs: Optional[str]
    ) -> Tuple[str, str]:
        """
        Generates both a quick patient-friendly overview and a detailed findings summary.
        """
        abnormal_params = [p for p in lab_parameters if p["status"] in ["LOW", "HIGH"]]
        normal_params = [p for p in lab_parameters if p["status"] == "NORMAL"]

        # Quick summary
        if not abnormal_params:
            quick = f"All tested laboratory parameters in your {doc_type} are within the standard reference intervals stated in your report."
        else:
            names = ", ".join([f"{p['parameter_name']} ({p['status'].lower()})" for p in abnormal_params])
            quick = f"Your {doc_type} shows {len(abnormal_params)} value(s) outside reference ranges: {names}. {len(normal_params)} other parameter(s) are normal."

        # Detailed summary
        details = []
        details.append(f"**Document Type:** {doc_type}")
        details.append(f"**Tests Performed:** {len(lab_parameters)} parameters analyzed.")

        if abnormal_params:
            details.append("\n**Important Findings / Out-of-Range Values:**")
            for p in abnormal_params:
                details.append(f"- **{p['parameter_name']}:** {p['result_value']} {p['unit']} (Status: **{p['status']}** | Reference: {p['reference_range']}) — {p['interpretation']}")

        if normal_params:
            details.append("\n**Normal Values:**")
            for p in normal_params:
                details.append(f"- **{p['parameter_name']}:** {p['result_value']} {p['unit']} (Reference: {p['reference_range']})")

        if doctor_obs:
            details.append(f"\n**Doctor Observations:** {doctor_obs}")
        elif findings:
            details.append(f"\n**Report Conclusion:** {findings}")

        return quick, "\n".join(details)
