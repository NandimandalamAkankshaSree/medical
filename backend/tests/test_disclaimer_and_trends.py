import os
import io
import json
import urllib.request
import urllib.parse
from docx import Document

BASE_URL = "http://127.0.0.1:8000/api"

def create_sample_docx_report(file_path: str, patient_name: str, patient_id: str):
    doc = Document()
    doc.add_heading("Apex Diagnostic Care - Clinical Laboratory Report", 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Patient Name: {patient_name}\n")
    p.add_run(f"Patient ID: {patient_id}\n")
    p.add_run("Age/Sex: 45 years / Male\n")
    p.add_run("Doctor: Dr. Sarah Jenkins (MD Nephrology)\n")
    p.add_run("Date of Report: 2026-08-27\n")
    
    doc.add_heading("Kidney Function Panel", level=1)
    
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test Parameter'
    hdr_cells[1].text = 'Result'
    hdr_cells[2].text = 'Unit'
    hdr_cells[3].text = 'Reference Range'
    
    tests = [
        ("Serum Creatinine", "1.8", "mg/dL", "0.6 - 1.2"),
        ("Potassium", "5.4", "mEq/L", "3.5 - 5.0"),
        ("Blood Urea Nitrogen", "28", "mg/dL", "7 - 20")
    ]
    
    for test, res, unit, ref in tests:
        row_cells = table.add_row().cells
        row_cells[0].text = test
        row_cells[1].text = res
        row_cells[2].text = unit
        row_cells[3].text = ref
        
    doc.add_paragraph("Clinical Impression: Renal function test indicates mild elevation in creatinine and BUN.")
    doc.save(file_path)

def create_multipart_form(fields: dict, files: dict):
    boundary = "----WebKitFormBoundary7MA4YWxkTrZu0gW"
    body = io.BytesIO()
    
    for k, v in fields.items():
        body.write(f"--{boundary}\r\n".encode())
        body.write(f'Content-Disposition: form-data; name="{k}"\r\n\r\n'.encode())
        body.write(f"{v}\r\n".encode())
        
    for k, (filename, filedata, content_type) in files.items():
        body.write(f"--{boundary}\r\n".encode())
        body.write(f'Content-Disposition: form-data; name="{k}"; filename="{filename}"\r\n'.encode())
        body.write(f"Content-Type: {content_type}\r\n\r\n".encode())
        body.write(filedata)
        body.write(b"\r\n")
        
    body.write(f"--{boundary}--\r\n".encode())
    return body.getvalue(), f"multipart/form-data; boundary={boundary}"

def test_disclaimer_and_trends():
    print("=== TESTING EXTERNAL REPORT ACCEPTANCE WITH DISCLAIMER & TRENDS ===")

    logged_in_patient_id = "user_11"
    test_doc_path = "tests/test_report_external_ramesh.docx"
    doc_id = None

    try:
        # 1. Upload external person's report (Ramesh Kumar)
        print("\n--- 1. Uploading External Report (Ramesh Kumar) for user_11 ---")
        create_sample_docx_report(test_doc_path, "Ramesh Kumar", "patient_ext_01")

        with open(test_doc_path, "rb") as f:
            doc_bytes = f.read()

        fields = {
            "patient_id": logged_in_patient_id,
            "document_type": "Kidney Report",
            "report_date": "2026-08-27"
        }
        files = {
            "file": ("test_report_external_ramesh.docx", doc_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }

        body, content_type = create_multipart_form(fields, files)
        req = urllib.request.Request(
            f"{BASE_URL}/documents/upload",
            data=body,
            headers={"Content-Type": content_type},
            method="POST"
        )

        with urllib.request.urlopen(req) as res:
            data = json.loads(res.read().decode("utf-8"))
            doc_id = data.get("document_id")
            is_mismatch = data.get("is_name_mismatch")
            disclaimer = data.get("disclaimer_note")
            extracted_name = data.get("patient_name_extracted")
            params = data.get("parameters_extracted")

            print(f"Uploaded Doc ID: {doc_id} | Parameters Extracted: {params}")
            print(f"is_name_mismatch: {is_mismatch}")
            print(f"patient_name_extracted: '{extracted_name}'")
            clean_disclaimer = (disclaimer or "").encode('ascii', 'ignore').decode('ascii')
            print(f"disclaimer_note: {clean_disclaimer}")

            assert is_mismatch is True, "Document must be flagged with is_name_mismatch=True"
            assert "ramesh kumar" in (extracted_name or "").lower(), "Must capture extracted name 'Ramesh Kumar'"
            assert "disclaimer" in (disclaimer or "").lower(), "Must have disclaimer note"
            print(" -> PASS: External report accepted and stamped with identity disclaimer!")

        # 2. Check Document List includes disclaimer
        print("\n--- 2. Checking Patient Document List Endpoint ---")
        req_docs = urllib.request.Request(f"{BASE_URL}/documents/patient/{logged_in_patient_id}")
        with urllib.request.urlopen(req_docs) as res:
            docs_list = json.loads(res.read().decode("utf-8"))
            target_doc = next((d for d in docs_list if d["id"] == doc_id), None)
            assert target_doc is not None, "Uploaded doc must be in documents list"
            assert target_doc.get("is_name_mismatch") is True
            clean_list_disclaimer = (target_doc.get('disclaimer_note') or "").encode('ascii', 'ignore').decode('ascii')
            print(f" -> Found in document list with mismatch flag: {clean_list_disclaimer}")
            print(" -> PASS: Document list includes disclaimer and mismatch metadata!")

        # 3. Check Health Trends Endpoint includes these parameters & disclaimer
        print("\n--- 3. Checking Health Trends & Longitudinal Graph Endpoint ---")
        req_trends = urllib.request.Request(f"{BASE_URL}/visualization/trends/{logged_in_patient_id}")
        with urllib.request.urlopen(req_trends) as res:
            trends_data = json.loads(res.read().decode("utf-8"))
            disclaimers = trends_data.get("disclaimers", [])
            print(f"Total Reports in Trends: {trends_data.get('total_reports')}")
            print(f"Disclaimers in Trends: {len(disclaimers)}")
            for d in disclaimers:
                clean_trend_disclaimer = (d.get('disclaimer_note') or "").encode('ascii', 'ignore').decode('ascii')
                print(f" -> Trend Disclaimer: {clean_trend_disclaimer}")

            assert len(disclaimers) > 0, "Trends response must include disclaimers for external reports"
            assert trends_data.get("has_external_reports") is True

            # Verify parameter is present in trends series
            creat_series = next((t for t in trends_data.get("trends", []) if "creatinine" in t["parameter_name"].lower()), None)
            assert creat_series is not None, "Creatinine series must exist in trends"
            print(f"Creatinine Data Points Count: {len(creat_series['data_points'])}")
            print(" -> PASS: Health trends includes external report data points with disclaimer!")

        # 4. Check Chat Assistant includes disclaimer when grounded in external document
        print("\n--- 4. Checking AI Chat Assistant with External Document ---")
        chat_req_payload = {
            "patient_id": logged_in_patient_id,
            "document_id": doc_id,
            "message": "What is my creatinine value in this report?"
        }
        req_chat = urllib.request.Request(
            f"{BASE_URL}/chat",
            data=json.dumps(chat_req_payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        with urllib.request.urlopen(req_chat) as res:
            chat_res = json.loads(res.read().decode("utf-8"))
            content = chat_res.get("content", "")
            clean_content = content.encode("ascii", "ignore").decode("ascii")
            print(f"Chat Response Snippet:\n{clean_content[:250]}...\n")
            assert "disclaimer" in content.lower() or "not your personal" in content.lower() or "ramesh kumar" in content.lower(), "Chat response must include identity disclaimer note!"
            print(" -> PASS: AI Chat Assistant attaches clear identity disclaimer notice!")

    finally:
        # Clean up test document
        if doc_id:
            del_req = urllib.request.Request(f"{BASE_URL}/documents/{doc_id}", method="DELETE")
            with urllib.request.urlopen(del_req) as del_res:
                print(f"\n -> Cleaned up test document {doc_id}.")
        if os.path.exists(test_doc_path):
            os.remove(test_doc_path)

    print("\n=======================================================")
    print(" EXTERNAL REPORT ACCEPTANCE & TRENDS PASSED 100%!")
    print("=======================================================")

if __name__ == "__main__":
    test_disclaimer_and_trends()
