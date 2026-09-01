import os
import io
import json
import urllib.request
import urllib.parse
from docx import Document

BASE_URL = "http://127.0.0.1:8000/api"

def create_sample_docx_report(file_path: str, patient_name: str, patient_id: str):
    doc = Document()
    doc.add_heading("CityCare Multispeciality Hospital - Clinical Laboratory Report", 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Patient Name: {patient_name}\n")
    p.add_run(f"Patient ID: {patient_id}\n")
    p.add_run("Age/Sex: 21 years / Female\n")
    p.add_run("Doctor: Dr. Sarah Jenkins (MD Nephrology)\n")
    p.add_run("Date of Report: 2026-08-26\n")
    
    doc.add_heading("Complete Blood & Kidney Panel", level=1)
    
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test Parameter'
    hdr_cells[1].text = 'Result'
    hdr_cells[2].text = 'Unit'
    hdr_cells[3].text = 'Reference Range'
    
    tests = [
        ("Serum Creatinine", "0.9", "mg/dL", "0.6 - 1.2"),
        ("Potassium", "4.2", "mEq/L", "3.5 - 5.0"),
        ("Hemoglobin", "13.5", "g/dL", "12.0 - 15.5"),
        ("eGFR", "105", "mL/min/1.73m2", "> 90")
    ]
    
    for test, res, unit, ref in tests:
        row_cells = table.add_row().cells
        row_cells[0].text = test
        row_cells[1].text = res
        row_cells[2].text = unit
        row_cells[3].text = ref
        
    doc.add_paragraph("Clinical Impression: All renal and metabolic parameters are within normal physiological limits.")
    doc.save(file_path)

def create_multipart_form(fields: dict, files: dict):
    boundary = "----WebKitFormBoundary7MA4YWxkTrZu0gW"
    body = io.BytesIO()
    
    for k, v in fields.items():
        body.write(f"--{boundary}\r\n".encode())
        body.write(f'Content-Disposition: form-data; name="{k}"\r\n\r\n'.encode())
        body.write(f"{v}\r\n".encode())
        
    for k, (filename, filedata, content_type) in files.items():
        body.write(f"--{boundary}\r\n".encode())
        body.write(f'Content-Disposition: form-data; name="{k}"; filename="{filename}"\r\n'.encode())
        body.write(f"Content-Type: {content_type}\r\n\r\n".encode())
        body.write(filedata)
        body.write(b"\r\n")
        
    body.write(f"--{boundary}--\r\n".encode())
    return body.getvalue(), f"multipart/form-data; boundary={boundary}"

def test_patient_name_guardrail():
    print("=== TESTING PATIENT NAME IDENTITY & OWNERSHIP GUARDRAIL ===")

    logged_in_patient_id = "user_11"
    # user_11 full_name is "Nandimandalam Akanksha Sree"

    other_person_doc_path = "tests/test_report_ramesh_kumar.docx"
    matching_person_doc_path = "tests/test_report_akanksha_sree.docx"

    try:
        # 1. Test uploading report with DIFFERENT person's name (Ramesh Kumar)
        print("\n--- 1. Testing Upload of Another Person's Report (Ramesh Kumar) ---")
        create_sample_docx_report(other_person_doc_path, "Ramesh Kumar", "patient_999")
        
        with open(other_person_doc_path, "rb") as f:
            doc_bytes = f.read()

        fields = {
            "patient_id": logged_in_patient_id,
            "document_type": "Kidney Report",
            "report_date": "2026-08-26"
        }
        files = {
            "file": ("test_report_ramesh_kumar.docx", doc_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }

        body, content_type = create_multipart_form(fields, files)
        req = urllib.request.Request(
            f"{BASE_URL}/documents/upload",
            data=body,
            headers={"Content-Type": content_type},
            method="POST"
        )

        try:
            with urllib.request.urlopen(req) as res:
                assert False, "Upload of another person's report MUST be rejected with HTTP 400!"
        except urllib.error.HTTPError as e:
            err_body = json.loads(e.read().decode("utf-8"))
            detail = err_body.get("detail", "")
            print(f"[HTTP {e.code}] Server Response Detail:\n{detail}")
            assert e.code == 400, f"Expected 400 Bad Request, got {e.code}"
            assert "upload your own" in detail.lower() or "does not match" in detail.lower(), "Must state 'Please upload your own medical reports'!"
            assert "ramesh kumar" in detail.lower(), "Must cite the mismatched name from the report!"
            print(" -> PASS: Correctly rejected with 'Please upload your own medical reports' error!")

        # 2. Test uploading report with MATCHING person's name (Akanksha Sree)
        print("\n--- 2. Testing Upload of Matching Patient Report (Nandimandalam Akanksha Sree) ---")
        create_sample_docx_report(matching_person_doc_path, "Nandimandalam Akanksha Sree", logged_in_patient_id)

        with open(matching_person_doc_path, "rb") as f:
            doc_bytes_match = f.read()

        files_match = {
            "file": ("test_report_akanksha_sree.docx", doc_bytes_match, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }
        body_match, content_type_match = create_multipart_form(fields, files_match)
        req_match = urllib.request.Request(
            f"{BASE_URL}/documents/upload",
            data=body_match,
            headers={"Content-Type": content_type_match},
            method="POST"
        )

        with urllib.request.urlopen(req_match) as res_match:
            data = json.loads(res_match.read().decode("utf-8"))
            doc_id = data.get("document_id")
            params = data.get("parameters_extracted")
            print(f"Uploaded Doc ID: {doc_id} | Parameters Extracted: {params}")
            assert params > 0, "Valid matching patient report must extract parameters!"
            print(" -> PASS: Matching patient report successfully accepted and indexed!")

            # Clean up uploaded test doc
            del_req = urllib.request.Request(f"{BASE_URL}/documents/{doc_id}", method="DELETE")
            with urllib.request.urlopen(del_req) as del_res:
                print(f" -> Cleaned up test document {doc_id}.")

    finally:
        for p in [other_person_doc_path, matching_person_doc_path]:
            if os.path.exists(p):
                os.remove(p)

    print("\n=======================================================")
    print(" PATIENT NAME OWNERSHIP GUARDRAIL PASSED 100%!")
    print("=======================================================")

if __name__ == "__main__":
    test_patient_name_guardrail()
